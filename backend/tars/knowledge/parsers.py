"""知识库文档解析器。"""
from __future__ import annotations

import csv
import os
import re


def _read_text_file(file_path: str) -> str:
    """Try UTF-8 first, then common Chinese legacy encodings."""
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk"):
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


class TextParser:
    def parse(self, file_path: str) -> str:
        return _read_text_file(file_path)


class MarkdownParser:
    def parse(self, file_path: str) -> str:
        text = _read_text_file(file_path)
        # 保留结构信息，去掉大部分 markdown 标记。
        text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
        text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"`{1,3}", "", text)
        text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
        return text.strip()


class PDFParser:
    def parse(self, file_path: str) -> str:
        import fitz

        texts = []
        with fitz.open(file_path) as doc:
            for page in doc:
                texts.append(page.get_text("text"))
        return "\n".join(t for t in texts if t).strip()


class DOCXParser:
    """Word 文档解析器。

    先尝试 python-docx（支持表格），失败时回退到 zipfile + XML 直接提取。
    """

    def parse(self, file_path: str) -> str:
        file_size = os.path.getsize(file_path)
        if file_size > 200 * 1024 * 1024:
            raise ValueError(f"文档过大（{file_size / 1024 / 1024:.0f}MB），超过 200MB 上限")

        # 优先用 python-docx（完整支持表格）
        try:
            return self._parse_with_python_docx(file_path)
        except Exception as e1:
            # 回退：直接用 zipfile + XML 提取（兼容损坏/复杂文档）
            try:
                return self._parse_with_xml_fallback(file_path)
            except Exception as e2:
                raise ValueError(
                    f"Word 文档解析失败。python-docx: {e1}；xml-fallback: {e2}"
                ) from e2

    def _parse_with_python_docx(self, file_path: str) -> str:
        from docx import Document

        document = Document(file_path)
        parts: list[str] = []

        # 1. 提取段落（含标题层级标记）
        for p in document.paragraphs:
            text = (p.text or "").strip()
            if not text:
                continue
            style_name = (p.style.name or "").lower()
            if "heading" in style_name or "标题" in style_name:
                level = 1
                for ch in style_name:
                    if ch.isdigit():
                        level = int(ch)
                        break
                prefix = "#" * min(level, 6)
                parts.append(f"\n{prefix} {text}")
            else:
                parts.append(text)

        # 2. 提取表格内容
        for table in document.tables:
            parts.append("")  # 表格前空行
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts).strip()

    def _parse_with_xml_fallback(self, file_path: str) -> str:
        """直接用 zipfile 读 document.xml，兼容 python-docx 无法处理的文档。"""
        import zipfile
        import xml.etree.ElementTree as ET

        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        with zipfile.ZipFile(file_path) as z:
            doc_xml = z.read("word/document.xml")
            # 读取 styles.xml 构建标题样式 → 级别映射
            heading_style_levels: dict[str, int] = {}
            try:
                styles_xml = z.read("word/styles.xml")
                styles_root = ET.fromstring(styles_xml)
                for style in styles_root.iter(f"{{{ns}}}style"):
                    style_id = style.get(f"{{{ns}}}styleId", "")
                    name_el = style.find(f"{{{ns}}}name")
                    name = (name_el.get(f"{{{ns}}}val") or "").lower() if name_el is not None else ""
                    # 匹配 "heading N" 或 "标题 N" 样式的名称
                    if "heading" in name or "标题" in name:
                        level = 1
                        for ch in name:
                            if ch.isdigit():
                                level = int(ch)
                                break
                        heading_style_levels[style_id] = level
                    # 也检查 outlineLvl
                    pPr = style.find(f"{{{ns}}}pPr")
                    if pPr is not None:
                        ol = pPr.find(f"{{{ns}}}outlineLvl")
                        if ol is not None:
                            lv = int(ol.get(f"{{{ns}}}val", "0")) + 1
                            if style_id not in heading_style_levels:
                                heading_style_levels[style_id] = lv
            except Exception:
                heading_style_levels = {}

        root = ET.fromstring(doc_xml)
        parts: list[str] = []

        for p in root.iter(f"{{{ns}}}p"):
            # 检查段落样式判断是否为标题
            pPr = p.find(f"{{{ns}}}pPr")
            is_heading = False
            heading_level = 0
            if pPr is not None:
                pStyle = pPr.find(f"{{{ns}}}pStyle")
                if pStyle is not None:
                    style_val = pStyle.get(f"{{{ns}}}val") or ""
                    # 先用 styles.xml 映射查找
                    if style_val in heading_style_levels:
                        is_heading = True
                        heading_level = heading_style_levels[style_val]
                    # 回退：检查样式名含 heading/标题
                    elif "heading" in style_val.lower() or "标题" in style_val:
                        is_heading = True
                        for ch in style_val:
                            if ch.isdigit():
                                heading_level = int(ch)
                                break
                # 某些文档用 outlineLvl
                if not is_heading:
                    ol = pPr.find(f"{{{ns}}}outlineLvl")
                    if ol is not None:
                        heading_level = int(ol.get(f"{{{ns}}}val", "0")) + 1
                        is_heading = True

            # 收集文本
            texts = []
            for t in p.iter(f"{{{ns}}}t"):
                if t.text:
                    texts.append(t.text)
            text = "".join(texts).strip()

            if not text:
                continue

            if is_heading and heading_level > 0:
                prefix = "#" * min(heading_level, 6)
                parts.append(f"\n{prefix} {text}")
            else:
                parts.append(text)

        # 也提取表格
        for tbl in root.iter(f"{{{ns}}}tbl"):
            parts.append("")
            for row in tbl.iter(f"{{{ns}}}tr"):
                cells = []
                for tc in row.iter(f"{{{ns}}}tc"):
                    cell_texts = []
                    for t in tc.iter(f"{{{ns}}}t"):
                        if t.text:
                            cell_texts.append(t.text)
                    cell_text = "".join(cell_texts).strip()
                    if cell_text:
                        cells.append(cell_text)
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts).strip()


class CSVParser:
    def parse(self, file_path: str) -> str:
        text = _read_text_file(file_path)
        try:
            from io import StringIO

            reader = csv.reader(StringIO(text))
            rows = [", ".join(cell.strip() for cell in row if cell is not None) for row in reader]
            return "\n".join(r for r in rows if r).strip()
        except Exception:
            return text.strip()


class XLSXParser:
    def parse(self, file_path: str) -> str:
        try:
            from openpyxl import load_workbook
        except ImportError as e:
            raise ValueError("解析 .xlsx 需要安装 openpyxl") from e

        wb = load_workbook(file_path, read_only=True, data_only=True)
        lines: list[str] = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    lines.append(", ".join(cells))
        wb.close()
        return "\n".join(lines).strip()


class DocumentParser:
    """按文件扩展名选择解析器。"""

    def __init__(self):
        self._parsers = {
            ".txt": TextParser(),
            ".md": MarkdownParser(),
            ".pdf": PDFParser(),
            ".docx": DOCXParser(),
            ".csv": CSVParser(),
            ".xlsx": XLSXParser(),
        }

    def parse(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        parser = self._parsers.get(ext)
        if not parser:
            raise ValueError(f"暂不支持的文档类型: {ext or 'unknown'}")
        return parser.parse(file_path)
